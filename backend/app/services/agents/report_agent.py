import csv
import uuid
from pathlib import Path
from typing import Literal

from docx import Document as DocxDocument
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.report import ReportModel

REPORT_FORMATS = ("csv", "xlsx", "docx", "pdf")


class ReportAgentError(Exception):
    pass


def _reports_dir() -> Path:
    path = Path(settings.report_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _write_csv(path: Path, columns: list[str], rows: list[list]) -> None:
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(columns)
        writer.writerows(rows)


def _write_xlsx(path: Path, title: str, columns: list[str], rows: list[list]) -> None:
    # Minimal "company template": bold header row, frozen header, sized columns.
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    ws.append(columns)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"
    for i, col_name in enumerate(columns, start=1):
        width = max(len(str(col_name)), *(len(str(r[i - 1])) for r in rows)) if rows else len(str(col_name))
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = min(width + 2, 60)
    wb.save(path)


def _write_docx(path: Path, title: str, columns: list[str], rows: list[list]) -> None:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for cell, col_name in zip(table.rows[0].cells, columns):
        cell.text = str(col_name)
        cell.paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
    doc.save(path)


def _write_pdf(path: Path, title: str, columns: list[str], rows: list[list]) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    elements = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    table = Table([columns] + [[str(v) for v in row] for row in rows], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F6BFF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F5FA")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)
    doc.build(elements)


_WRITERS = {
    "csv": lambda path, title, columns, rows: _write_csv(path, columns, rows),
    "xlsx": _write_xlsx,
    "docx": _write_docx,
    "pdf": _write_pdf,
}


def generate_report(
    db: Session,
    *,
    title: str,
    fmt: Literal["csv", "xlsx", "docx", "pdf"],
    columns: list[str],
    rows: list[list],
    owner_id: uuid.UUID | None = None,
    department: str | None = None,
) -> ReportModel:
    if fmt not in REPORT_FORMATS:
        raise ReportAgentError(f"Unsupported report format: {fmt}")
    if not columns:
        raise ReportAgentError("A report needs at least one column")
    for row in rows:
        if len(row) != len(columns):
            raise ReportAgentError("Every row must have the same number of values as there are columns")

    report_id = uuid.uuid4()
    filename = f"{report_id}.{fmt}"
    path = _reports_dir() / filename

    _WRITERS[fmt](path, title, columns, rows)

    row_model = ReportModel(
        id=report_id, title=title, format=fmt, file_path=str(path), row_count=len(rows),
        owner_id=owner_id, department=department,
    )
    db.add(row_model)
    db.commit()
    db.refresh(row_model)
    return row_model
